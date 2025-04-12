import streamlit as st
import google.generativeai as genai
from google.generativeai.types import GenerateContentResponse
import pdfplumber
from fpdf import FPDF
from docx import Document
from io import BytesIO
import matplotlib.pyplot as plt
import time

# Must be first
st.set_page_config(page_title="EvalBuddy", layout="wide")

# =================== CSS ===================
def apply_custom_css():
    st.markdown("""
    <style>
        .stApp { background-color: #1A1A1A; color: #F0F0F0; }
        h1, h2, h3, h4, h5, h6 { color: #F0F0F0 !important; }
        .stTextInput input, .stTextArea textarea, .stSelectbox, .stSlider {
            background-color: #2D2D2D !important; color: #F0F0F0 !important;
        }
        .stButton button {
            background-color: #FF7F50 !important;
            color: #FFFFFF !important;
            border-radius: 8px !important;
        }
        .stChatMessage[data-testid="stChatMessage-user"] { background-color: #2D3748 !important; }
        .stChatMessage[data-testid="stChatMessage-assistant"] { background-color: #333333 !important; }
        .processing-indicator {
            color: #FF7F50; padding: 10px;
            animation: pulse 1.5s infinite ease-in-out;
            font-weight: bold;
        }
        @keyframes pulse {
            0% { opacity: 0.6; }
            50% { opacity: 1; }
            100% { opacity: 0.6; }
        }
    </style>
    """, unsafe_allow_html=True)

# =================== PDF UPLOAD ===================
def pdf_upload_area():
    uploaded_pdf = st.file_uploader("📄 Upload Evaluation PDF", type=["pdf"])
    if uploaded_pdf:
        with pdfplumber.open(uploaded_pdf) as pdf:
            content = ""
            for page in pdf.pages:
                content += page.extract_text() or ""
        st.session_state.pdf_content = content
        with st.expander("🔍 Preview Extracted PDF Text"):
            st.text_area("PDF Content", content, height=300, disabled=True)

# =================== EXPORT HELPERS ===================
def export_chat_to_pdf(messages):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt="EvalBuddy Chat Export", ln=True, align="C")
    pdf.ln()
    for msg in messages:
        role = msg["role"].capitalize()
        content = msg["content"]
        pdf.multi_cell(0, 10, f"{role}: {content}")
        pdf.ln()
    output = BytesIO()
    pdf.output(output)
    output.seek(0)
    return output

def export_chat_to_docx(messages):
    doc = Document()
    doc.add_heading("EvalBuddy Chat Export", 0)
    for msg in messages:
        role = msg["role"].capitalize()
        content = msg["content"]
        doc.add_paragraph(f"{role}: {content}")
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

def export_logic_model_to_docx(data):
    doc = Document()
    doc.add_heading("Logic Model", 0)
    for section, value in data.items():
        doc.add_heading(section, level=1)
        doc.add_paragraph(value)
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

def export_chart_to_pdf(fig):
    buf = BytesIO()
    fig.savefig(buf, format="png", dpi=300)
    buf.seek(0)
    pdf = FPDF()
    pdf.add_page()
    pdf.image(buf, x=10, y=20, w=180)
    output = BytesIO()
    pdf.output(output)
    output.seek(0)
    return output

# =================== GEMINI SETUP ===================
@st.cache_resource
def initialize_chat_session(model_name, temperature):
    generation_config = {
        "temperature": temperature,
        "top_p": 0.95,
        "top_k": 40,
        "max_output_tokens": 8192,
    }
    model = genai.GenerativeModel(
        model_name=model_name,
        generation_config=generation_config,
    )
    return model.start_chat()

def stream_response(response: GenerateContentResponse):
    for chunk in response:
        yield chunk.text

def display_processing_indicator():
    return st.markdown('<div class="processing-indicator">EvalBuddy is thinking...</div>', unsafe_allow_html=True)

# =================== CHAT TAB ===================
def home_page():
    st.header("Chat with EvalBuddy")
    st.caption("EvalBuddy is an AI assistant for evaluation guidance.")
    pdf_upload_area()

    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

    user_input = st.chat_input("How can I help with your evaluation work today?")
    if user_input:
        current_message = {"role": "user", "content": user_input}
        st.session_state.messages.append(current_message)

        with st.chat_message("user"):
            st.markdown(current_message["content"])
        with st.chat_message("assistant"):
            placeholder = st.empty()
            indicator = display_processing_indicator()

            if "chat_session" not in st.session_state:
                st.session_state.chat_session = initialize_chat_session(
                    st.session_state.model_name,
                    st.session_state.temperature
                )
                st.session_state.chat_session.send_message(f"System: {st.session_state.system_prompt}")
                if st.session_state.pdf_content:
                    st.session_state.chat_session.send_message(
                        f"The following is the content of an uploaded PDF document:\n\n{st.session_state.pdf_content}"
                    )

            try:
                response = st.session_state.chat_session.send_message(current_message["content"], stream=True)
                full_response = ""
                for chunk in stream_response(response):
                    full_response += chunk
                    placeholder.markdown(full_response + "▌")
                    time.sleep(0.001)  # Reduced for speed
                indicator.empty()
                placeholder.markdown(full_response)
                st.session_state.messages.append({"role": "assistant", "content": full_response})
            except Exception as e:
                indicator.empty()
                st.error(f"Error: {e}")

    st.markdown("---")
    st.subheader("📤 Export Conversation")
    format_choice = st.selectbox("Choose format", ["PDF", "Word (DOCX)"])
    if st.button("Download Chat"):
        if format_choice == "PDF":
            pdf = export_chat_to_pdf(st.session_state.messages)
            st.download_button("📄 Download PDF", pdf, "EvalBuddy_Chat.pdf", mime="application/pdf")
        else:
            doc = export_chat_to_docx(st.session_state.messages)
            st.download_button("📝 Download Word", doc, "EvalBuddy_Chat.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# =================== RESOURCES TAB ===================
def resources_page():
    st.header("Evaluation Resources")
    context = st.text_area("Describe your evaluation context:")
    if st.button("Get Recommendations"):
        st.info("⚠️ Recommendation system not yet connected.")

# =================== TOOLS TAB ===================
def evaluation_tools_page():
    st.header("Evaluation Tools")

    st.subheader("🗂️ Project Manager")
    project_name = st.text_input("Project Name")
    if "projects" not in st.session_state:
        st.session_state.projects = {}

    if st.button("💾 Save Project"):
        if not project_name:
            st.warning("Please enter a project name.")
        else:
            st.session_state.projects[project_name] = {
                "chat": st.session_state.messages.copy(),
                "logic_model": {
                    "Inputs": st.session_state.get("lm_inputs", ""),
                    "Activities": st.session_state.get("lm_activities", ""),
                    "Outputs": st.session_state.get("lm_outputs", ""),
                    "Outcomes": st.session_state.get("lm_outcomes", ""),
                    "Impact": st.session_state.get("lm_impact", "")
                },
                "chart": {
                    "type": st.session_state.get("chart_type", "Bar"),
                    "x_data": st.session_state.get("x_data", ""),
                    "y_data": st.session_state.get("y_data", ""),
                    "title": st.session_state.get("chart_title", "")
                }
            }
            st.success(f"Project '{project_name}' saved!")

    if st.button("📂 Load Project"):
        if project_name not in st.session_state.projects:
            st.warning("Project not found.")
        else:
            project = st.session_state.projects[project_name]
            st.session_state.messages = project["chat"]
            lm = project["logic_model"]
            st.session_state.lm_inputs = lm["Inputs"]
            st.session_state.lm_activities = lm["Activities"]
            st.session_state.lm_outputs = lm["Outputs"]
            st.session_state.lm_outcomes = lm["Outcomes"]
            st.session_state.lm_impact = lm["Impact"]
            ch = project["chart"]
            st.session_state.chart_type = ch["type"]
            st.session_state.x_data = ch["x_data"]
            st.session_state.y_data = ch["y_data"]
            st.session_state.chart_title = ch["title"]
            st.success(f"Project '{project_name}' loaded!")

    st.markdown("---")
    tool = st.selectbox("Choose Tool", ["Logic Model Builder", "Chart Generator"])

    if tool == "Logic Model Builder":
        st.subheader("🧩 Logic Model Builder")
        inputs = st.text_area("Inputs", value=st.session_state.get("lm_inputs", ""), key="lm_inputs")
        activities = st.text_area("Activities", value=st.session_state.get("lm_activities", ""), key="lm_activities")
        outputs = st.text_area("Outputs", value=st.session_state.get("lm_outputs", ""), key="lm_outputs")
        outcomes = st.text_area("Outcomes", value=st.session_state.get("lm_outcomes", ""), key="lm_outcomes")
        impact = st.text_area("Impact", value=st.session_state.get("lm_impact", ""), key="lm_impact")
        logic_model_data = {
            "Inputs": inputs,
            "Activities": activities,
            "Outputs": outputs,
            "Outcomes": outcomes,
            "Impact": impact
        }
        if st.button("📤 Export Logic Model (Word)"):
            doc = export_logic_model_to_docx(logic_model_data)
            st.download_button("Download DOCX", doc, "Logic_Model.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    elif tool == "Chart Generator":
        st.subheader("📊 Chart Generator")
        chart_type = st.selectbox("Chart Type", ["Bar", "Line"], key="chart_type")
        x_data = st.text_input("X-axis values (comma-separated)", st.session_state.get("x_data", ""), key="x_data")
        y_data = st.text_input("Y-axis values (comma-separated)", st.session_state.get("y_data", ""), key="y_data")
        title = st.text_input("Chart Title", st.session_state.get("chart_title", ""), key="chart_title")
        if st.button("Generate Chart"):
            x = [i.strip() for i in x_data.split(",")]
            y = [float(i.strip()) for i in y_data.split(",")]
            fig, ax = plt.subplots()
            if chart_type == "Bar":
                ax.bar(x, y)
            else:
                ax.plot(x, y, marker="o")
            ax.set_title(title)
            st.pyplot(fig)
            if st.button("📤 Export Chart (PDF)"):
                chart_pdf = export_chart_to_pdf(fig)
                st.download_button("Download Chart PDF", chart_pdf, "Chart.pdf", mime="application/pdf")

# =================== MAIN ===================
def main():
    apply_custom_css()
    st.session_state.setdefault("messages", [])
    st.session_state.setdefault("model_name", "gemini-pro")
    st.session_state.setdefault("temperature", 0.7)
    st.session_state.setdefault("system_prompt", "You are EvalBuddy, an AI assistant specialized in program evaluation.")
    st.session_state.setdefault("pdf_content", "")

    st.title("EvalBuddy")
    st.caption("Your Evaluation Assistant for Smarter Impact")

    tabs = st.tabs(["💬 Chat", "📚 Resources", "🛠️ Tools"])
    with tabs[0]:
        home_page()
    with tabs[1]:
        resources_page()
    with tabs[2]:
        evaluation_tools_page()

if __name__ == "__main__":
    main()
